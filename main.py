"""
Document Generator - FACTURA Generator + SOMATII Batch Generator
Local desktop application (Tkinter) for generating Word documents from templates.

Requirements:
    pip install -r requirements.txt

Usage:
    python main.py
"""

import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import os
import sys
import re
import threading
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from copy import deepcopy

from docx.shared import Pt

try:
    from docx import Document
except ImportError:
    print("python-docx is not installed. Run: pip install -r requirements.txt")
    sys.exit(1)

from field_config import FACTURA_FIELDS, PLACEHOLDER_MAP, UPPERCASE_BOLD_FIELDS, UNDERLINE_FIELDS
import somatie_generator

TEMPLATE_PATH = "templates/template_contract.docx"
OUTPUT_DIR = "output"
FIELD_FONT_SIZES = {"nr": 12, "data": 12}


def validate_date(value: str) -> bool:
    return bool(re.match(r"^\d{2}\.\d{2}\.\d{4}$", value))


def validate_number(value: str) -> bool:
    return value.strip().isdigit()


class DocumentGeneratorApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Document Generator")
        self.geometry("820x900")
        self.resizable(False, False)
        self.mode_var = tk.StringVar(value="factura")
        self.entries = {}
        self.currency_var = tk.StringVar(value="RON")
        self.curs_row = None
        self._build_ui()

    def _build_ui(self):
        top_frame = ttk.Frame(self, padding=10)
        top_frame.pack(fill="x")
        ttk.Label(top_frame, text="Select document type:", font=("Segoe UI", 10, "bold")).pack(anchor="w")
        radio_frame = ttk.Frame(top_frame)
        radio_frame.pack(fill="x", pady=5)
        ttk.Radiobutton(radio_frame, text="Contract Generator", variable=self.mode_var, value="factura", command=self._render_mode).pack(side="left", padx=(0, 20))
        ttk.Radiobutton(radio_frame, text="SOMATII Generator", variable=self.mode_var, value="somatii", command=self._render_mode).pack(side="left")
        ttk.Separator(self, orient="horizontal").pack(fill="x", padx=10, pady=5)
        self.mode_container = ttk.Frame(self)
        self.mode_container.pack(fill="both", expand=True)
        self._render_mode()

    def _clear_mode_container(self):
        for widget in self.mode_container.winfo_children():
            widget.destroy()
        self.entries.clear()

    def _render_mode(self):
        self._clear_mode_container()
        if self.mode_var.get() == "factura":
            self._render_factura_ui()
        else:
            self._render_somatii_ui()

    def _render_factura_ui(self):
        container = ttk.Frame(self.mode_container)
        container.pack(fill="both", expand=True)
        canvas = tk.Canvas(container, borderwidth=0, highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scroll_frame = ttk.Frame(canvas)
        scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True, padx=(10, 0))
        scrollbar.pack(side="right", fill="y")

        row = 0
        group_titles = {
            "nume": "Date companie",
            "termen_livrare": "Livrare si plata",
            "profil_ales": "Detalii contract",
        }
        for field in FACTURA_FIELDS:
            if field["key"] == "curs_bnr":
                continue
            if field["key"] in group_titles:
                ttk.LabelFrame(scroll_frame, text=group_titles[field["key"]], padding=(8, 4)).grid(
                    row=row, column=0, columnspan=3, sticky="ew", padx=8, pady=(8, 2)
                )
                row += 1
            ttk.Label(scroll_frame, text=field["label"], width=24, anchor="w").grid(row=row, column=0, sticky="w", padx=10, pady=6)
            entry = ttk.Entry(scroll_frame, width=38)
            entry.grid(row=row, column=1, sticky="w", padx=5, pady=6)
            if field.get("type") == "date":
                ttk.Label(scroll_frame, text="DD.MM.YYYY", foreground="#888").grid(row=row, column=2, sticky="w")
            self.entries[field["key"]] = entry
            row += 1

            if field["key"] == "valoare_contract":
                currency_frame = ttk.Frame(scroll_frame)
                currency_frame.grid(row=row, column=2, sticky="w", padx=(10, 5), pady=6)
                ttk.Label(currency_frame, text="Moneda:").pack(side="left", padx=(0, 8))
                ttk.Radiobutton(currency_frame, text="RON", variable=self.currency_var, value="RON", command=self._toggle_curs_row).pack(side="left")
                ttk.Radiobutton(currency_frame, text="EUR", variable=self.currency_var, value="EUR", command=self._toggle_curs_row).pack(side="left", padx=(10, 0))

            if field["key"] == "diferenta_contract":
                self.curs_row = ttk.Frame(scroll_frame)
                self.curs_row.grid(row=row, column=0, columnspan=3, sticky="w")
                ttk.Label(self.curs_row, text="Curs BNR:", width=24, anchor="w").pack(side="left", padx=10, pady=6)
                curs_entry = ttk.Entry(self.curs_row, width=38)
                curs_entry.pack(side="left", padx=5, pady=6)
                self.entries["curs_bnr"] = curs_entry
                row += 1

        self.entries["data"].bind("<KeyRelease>", self._update_data_livrare)
        self.entries["termen_livrare"].bind("<KeyRelease>", self._update_data_livrare)
        self.entries["valoare_contract"].bind("<KeyRelease>", self._update_contract_amounts)
        self._toggle_curs_row()

        bottom_frame = ttk.Frame(self.mode_container, padding=10)
        bottom_frame.pack(fill="x")
        tk.Button(bottom_frame, text="GENERATE", font=("Segoe UI", 14, "bold"), bg="#2e7d32", fg="white", activebackground="#1b5e20", height=2, command=self.on_generate_factura).pack(fill="x")
        self.factura_status_label = ttk.Label(bottom_frame, text="", foreground="#555")
        self.factura_status_label.pack(pady=(6, 0))

    def _collect_factura_values(self):
        values = {}
        for field in FACTURA_FIELDS:
            key = field["key"]
            raw_value = self.entries[key].get().strip()
            if field.get("required", True) and not raw_value:
                raise ValueError(f"Field '{field['label']}' is required.")
            if key == "curs_bnr" and self.currency_var.get() == "EUR" and not raw_value:
                raise ValueError("Field 'Curs BNR:' is required for EUR contracts.")
            if field.get("type") == "date" and raw_value and not validate_date(raw_value):
                raise ValueError(f"Field '{field['label']}' must be in DD.MM.YYYY format.")
            if field.get("type") == "number" and raw_value and not validate_number(raw_value):
                raise ValueError(f"Field '{field['label']}' must contain only digits.")
            values[key] = raw_value.upper() if key in UPPERCASE_BOLD_FIELDS else raw_value
        values = {key: value or "-" for key, value in values.items()}
        values["valoare_contract"] = self._with_currency(values["valoare_contract"])
        values["avans_contract"] = self._with_currency(values["avans_contract"])
        values["diferenta_contract"] = self._with_currency(values["diferenta_contract"])
        return values

    @staticmethod
    def _parse_amount(value):
        normalized = value.strip().replace(" ", "")
        if "," in normalized:
            normalized = normalized.replace(".", "").replace(",", ".")
        try:
            return Decimal(normalized)
        except (InvalidOperation, ValueError):
            return None

    def _with_currency(self, value):
        if value == "-":
            return value
        return f"{value} {self.currency_var.get()}"

    def _toggle_curs_row(self):
        if self.curs_row is None:
            return
        if self.currency_var.get() == "EUR":
            self.curs_row.grid()
        else:
            self.curs_row.grid_remove()
            self.entries["curs_bnr"].delete(0, "end")

    def _update_contract_amounts(self, _event=None):
        amount = self._parse_amount(self.entries["valoare_contract"].get())
        if amount is None:
            return
        advance = amount / 2
        difference = amount - advance
        self.entries["avans_contract"].delete(0, "end")
        self.entries["avans_contract"].insert(0, f"{advance:g}")
        self.entries["diferenta_contract"].delete(0, "end")
        self.entries["diferenta_contract"].insert(0, f"{difference:g}")

    def _update_data_livrare(self, _event=None):
        start_text = self.entries["data"].get().strip()
        days_text = self.entries["termen_livrare"].get().strip()
        try:
            current = datetime.strptime(start_text, "%d.%m.%Y").date()
            working_days = int(days_text)
        except (ValueError, TypeError):
            return
        if working_days < 0:
            return
        while working_days:
            current += timedelta(days=1)
            if current.weekday() < 5:
                working_days -= 1
        delivery_entry = self.entries["data_livrare"]
        delivery_entry.delete(0, "end")
        delivery_entry.insert(0, current.strftime("%d.%m.%Y"))

    def on_generate_factura(self):
        try:
            values = self._collect_factura_values()
        except ValueError as e:
            messagebox.showerror("Invalid input", str(e))
            return
        if not os.path.exists(TEMPLATE_PATH):
            messagebox.showerror("Template missing", f"Could not find template at:\n{TEMPLATE_PATH}\n\nPlace template_contract.docx there before generating.")
            return
        try:
            output_path = self.generate_factura_document(values)
        except Exception as e:
            messagebox.showerror("Generation failed", str(e))
            return
        self.factura_status_label.config(text=f"Saved: {output_path}")
        messagebox.showinfo("Success", f"Document generated:\n{output_path}")

    def generate_factura_document(self, values: dict) -> str:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        doc = Document(TEMPLATE_PATH)
        ordered_keys = sorted(values.keys(), key=lambda k: -len(PLACEHOLDER_MAP.get(k, "")))
        for key in ordered_keys:
            placeholder = PLACEHOLDER_MAP.get(key)
            if placeholder:
                replace_placeholder_everywhere(
                    doc, placeholder, values[key], key not in UNDERLINE_FIELDS,
                    key in UNDERLINE_FIELDS, FIELD_FONT_SIZES.get(key, 9)
                )
        safe_name = re.sub(r"[^A-Za-z0-9_\-]+", "_", values.get("nume", "Contract")).strip("_") or "Contract"
        output_path = os.path.join(OUTPUT_DIR, f"{safe_name}_CT-Vanzare-Cumparare-NICOMY-NDY.docx")
        doc.save(output_path)
        return output_path

    def _render_somatii_ui(self):
        info_frame = ttk.Frame(self.mode_container, padding=10)
        info_frame.pack(fill="x")
        info_text = (
            "Reads source_excel.xlsx (last sheet, column L = 'somam'), and generates:\n"
            "  - one somatie_<CompanyName>.docx per matching company\n"
            "  - one combined envelope_toate_companiile.docx (one page per company)\n"
            "  - one combined borderou_toate_companiile.docx (one table row per company)\n\n"
            "Required files in this folder: source_excel.xlsx, template_somatie.docx,\n"
            "template_envelope-2.docx, template_borderou-3.docx"
        )
        ttk.Label(info_frame, text=info_text, foreground="#444", justify="left").pack(anchor="w")
        btn_frame = ttk.Frame(self.mode_container, padding=10)
        btn_frame.pack(fill="x")
        self.somatii_generate_btn = tk.Button(btn_frame, text="GENERATE", font=("Segoe UI", 14, "bold"), bg="#2e7d32", fg="white", activebackground="#1b5e20", height=2, command=self.on_generate_somatii)
        self.somatii_generate_btn.pack(fill="x")
        log_frame = ttk.Frame(self.mode_container, padding=(10, 5, 10, 10))
        log_frame.pack(fill="both", expand=True)
        ttk.Label(log_frame, text="Output log:", font=("Segoe UI", 9, "bold")).pack(anchor="w")
        self.somatii_log = scrolledtext.ScrolledText(log_frame, height=24, font=("Consolas", 9), state="disabled", bg="#1e1e1e", fg="#d4d4d4", insertbackground="#d4d4d4")
        self.somatii_log.pack(fill="both", expand=True, pady=(4, 0))

    def _log_somatii(self, message: str):
        def append():
            self.somatii_log.config(state="normal")
            self.somatii_log.insert("end", message + "\n")
            self.somatii_log.see("end")
            self.somatii_log.config(state="disabled")
        self.after(0, append)

    def on_generate_somatii(self):
        self.somatii_log.config(state="normal")
        self.somatii_log.delete("1.0", "end")
        self.somatii_log.config(state="disabled")
        self.somatii_generate_btn.config(state="disabled", text="GENERATING...")
        threading.Thread(target=self._run_somatii_batch, daemon=True).start()

    def _run_somatii_batch(self):
        try:
            self._log_somatii(f"=== Starting batch run: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ===")
            generated = somatie_generator.run_batch(log=self._log_somatii)
            self._log_somatii(f"=== Finished. {len(generated)} file(s) generated. ===")
            self.after(0, lambda: messagebox.showinfo("Success", f"Generated {len(generated)} file(s). See log for details."))
        except Exception as e:
            self._log_somatii(f"ERROR: {e}")
            self.after(0, lambda: messagebox.showerror("Generation failed", str(e)))
        finally:
            self.after(0, lambda: self.somatii_generate_btn.config(state="normal", text="GENERATE"))


def replace_placeholder_everywhere(doc, placeholder: str, value: str, bold: bool = False, underline: bool = False, font_size: int = 9):
    targets = list(doc.paragraphs)
    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:
                targets.extend(cell.paragraphs)
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            targets.extend(header_footer.paragraphs)
            for table in header_footer.tables:
                for row_ in table.rows:
                    for cell in row_.cells:
                        targets.extend(cell.paragraphs)
    for paragraph in targets:
        _replace_in_paragraph(paragraph, placeholder, value, bold, underline, font_size)


def _replace_in_paragraph(paragraph, placeholder: str, value: str, bold: bool, underline: bool, font_size: int = 9):
    runs = list(paragraph.runs)
    full_text = "".join(run.text for run in runs) or paragraph.text
    if placeholder not in full_text:
        return
    if not runs:
        runs = [paragraph.add_run(full_text)]

    start = full_text.index(placeholder)
    end = start + len(placeholder)
    paragraph.clear()
    position = 0
    inserted = False

    for original_run in runs:
        run_text = original_run.text
        run_start = position
        run_end = run_start + len(run_text)

        if run_start < start and run_end > run_start:
            prefix_end = min(run_end, start)
            _add_preserved_run(paragraph, run_text[:prefix_end - run_start], original_run)

        if not inserted and run_start <= start < run_end:
            value_run = paragraph.add_run(value)
            value_run.bold = bold
            value_run.underline = underline
            value_run.font.name = "Arial"
            value_run.font.size = Pt(font_size)
            inserted = True

        if run_end > end:
            suffix_start = max(end, run_start) - run_start
            _add_preserved_run(paragraph, run_text[suffix_start:], original_run)
        elif run_start >= end:
            _add_preserved_run(paragraph, run_text, original_run)
        position = run_end

    if not inserted:
        value_run = paragraph.add_run(value)
        value_run.bold = bold
        value_run.underline = underline
        value_run.font.name = "Arial"
        value_run.font.size = Pt(font_size)


def _add_preserved_run(paragraph, text, original_run):
    if not text:
        return
    new_run = paragraph.add_run(text)
    if original_run._r.rPr is not None:
        new_run._r.get_or_add_rPr().extend(deepcopy(list(original_run._r.rPr)))


if __name__ == "__main__":
    app = DocumentGeneratorApp()
    app.mainloop()
