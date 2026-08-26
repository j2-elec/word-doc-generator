"""
Template inspector helper.

Run this against your real .docx template BEFORE using the main app,
to see all existing text so you can decide where to insert {{PLACEHOLDER}}
tokens, and to confirm the tokens were typed correctly once you've edited it.

Usage:
    python inspect_template.py templates/factura_template.docx
"""

import sys
from docx import Document


def dump_paragraphs(doc, label):
    print(f"\n--- {label} ---")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"[{i}] {p.text}")


def dump_tables(doc, label):
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- {label} table {t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"  row {r_idx}, col {c_idx}: {cell.text}")


def main(path):
    doc = Document(path)
    dump_paragraphs(doc, "Body paragraphs")
    dump_tables(doc, "Body")

    for s_idx, section in enumerate(doc.sections):
        dump_paragraphs(section.header, f"Section {s_idx} header")
        dump_tables(section.header, f"Section {s_idx} header")
        dump_paragraphs(section.footer, f"Section {s_idx} footer")
        dump_tables(section.footer, f"Section {s_idx} footer")

    print("\nDone. Look for the spots where NR, Data, NAME, etc. should go,")
    print("then edit the .docx in Word and type the matching {{PLACEHOLDER}}")
    print("token exactly (see field_config.py for the full list) in that spot.")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python inspect_template.py <path-to-docx>")
        sys.exit(1)
    main(sys.argv[1])
