"""
SOMATII batch generator core logic.

Reads source_excel.xlsx (last sheet), filters column L == "somam"/"som\u0103m",
and for each matching row generates:
  - one somatie_<CompanyName>.docx (from template_somatie.docx)
  - appends a page to a single combined envelope document (from template_envelope-2.docx)
  - appends a row to a single combined borderou document (from template_borderou-3.docx)

All templates and outputs live in the same folder as this script (or the
folder the .exe is run from).

Column layout in source_excel.xlsx (last sheet):
    B -> CompanyName
    H -> CompanyTotal   (decimal, 2 decimals)
    I -> CompanyCUI
    J -> CompanyJ
    K -> CompanyAddress (raw, "JUD. X, LOCALITATEA, STRADA, NR. Y" format)
    L -> filter column, keep rows where value == "somam" (matches "som\u0103m"
         with or without diacritics, case-insensitive)
"""

import os
import re
import copy

import openpyxl
from docx import Document
from docx.table import _Row
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_BREAK

from address_utils import format_address_somatie, format_address_envelope, format_address_borderou

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

SOURCE_EXCEL = os.path.join(BASE_DIR, "source_excel.xlsx")
TEMPLATE_SOMATIE = os.path.join(BASE_DIR, "template_somatie.docx")
TEMPLATE_ENVELOPE = os.path.join(BASE_DIR, "template_envelope-2.docx")
TEMPLATE_BORDEROU = os.path.join(BASE_DIR, "template_borderou-3.docx")

FILTER_COLUMN = "L"
FILTER_VALUE_VARIANTS = {"somam", "som\u0103m"}

COL_NAME = "B"
COL_TOTAL = "H"
COL_CUI = "I"
COL_J = "J"
COL_ADDRESS = "K"


def _normalize(value) -> str:
    if value is None:
        return ""
    value = str(value).strip().lower()
    value = (value.replace("\u0103", "a").replace("\u00e2", "a")
                   .replace("\u00ee", "i").replace("\u0219", "s").replace("\u021b", "t"))
    return value


def read_source_rows(log=print):
    if not os.path.exists(SOURCE_EXCEL):
        raise FileNotFoundError(f"Could not find {SOURCE_EXCEL}")

    wb = openpyxl.load_workbook(SOURCE_EXCEL, data_only=True)
    sheet = wb.worksheets[-1]
    log(f"Reading sheet: '{sheet.title}' (last sheet in workbook)")

    rows = []
    for row_idx in range(2, sheet.max_row + 1):
        filter_val = sheet[f"{FILTER_COLUMN}{row_idx}"].value
        if _normalize(filter_val) not in FILTER_VALUE_VARIANTS:
            continue

        name = sheet[f"{COL_NAME}{row_idx}"].value
        total = sheet[f"{COL_TOTAL}{row_idx}"].value
        cui = sheet[f"{COL_CUI}{row_idx}"].value
        j_val = sheet[f"{COL_J}{row_idx}"].value
        address = sheet[f"{COL_ADDRESS}{row_idx}"].value

        if not name or not address:
            log(f"  Row {row_idx}: SKIPPED (missing CompanyName or CompanyAddress)")
            continue

        if isinstance(total, (int, float)):
            total_str = f"{total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        else:
            total_str = str(total).strip() if total is not None else ""

        rows.append({
            "row": row_idx,
            "name": str(name).strip(),
            "address_raw": str(address).strip(),
            "cui": str(cui).strip() if cui is not None else "",
            "j": str(j_val).strip() if j_val is not None else "",
            "total": total_str,
        })
        log(f"  Row {row_idx}: MATCHED -> {name}")

    log(f"Total matched companies: {len(rows)}")
    return rows


def _iter_all_paragraphs(doc):
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:
                paragraphs.extend(cell.paragraphs)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            paragraphs.extend(hf.paragraphs)
            for table in hf.tables:
                for row_ in table.rows:
                    for cell in row_.cells:
                        paragraphs.extend(cell.paragraphs)
    return paragraphs


def _replace_token_in_paragraph(paragraph, token, value):
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        full_text = paragraph.text
    if token not in full_text:
        return False
    new_text = full_text.replace(token, value)
    if paragraph.runs:
        bold = paragraph.runs[0].bold
        italic = paragraph.runs[0].italic
        underline = paragraph.runs[0].underline
        for run in list(paragraph.runs):
            run.text = ""
        paragraph.runs[0].text = new_text
        paragraph.runs[0].bold = bold
        paragraph.runs[0].italic = italic
        paragraph.runs[0].underline = underline
    else:
        paragraph.add_run(new_text)
    return True


def replace_tokens_in_doc(doc, mapping: dict):
    for paragraph in _iter_all_paragraphs(doc):
        for token, value in mapping.items():
            _replace_token_in_paragraph(paragraph, token, value)


# ----------------------------------------------------------------------
# SOMATIE (one .docx per company)
# ----------------------------------------------------------------------

def generate_somatie_docs(rows, out_dir, log=print):
    generated = []
    for company in rows:
        log(f"[SOMATIE] Reading company: {company['name']}")
        doc = Document(TEMPLATE_SOMATIE)
        address_formatted = format_address_somatie(company["address_raw"])
        mapping = {
            "CompanyName": company["name"],
            "CompanyAddress": address_formatted,
            "CompanyCUI": company["cui"],
            "CompanyJ": company["j"],
            "CompanyTotal": company["total"],
        }
        replace_tokens_in_doc(doc, mapping)

        safe_name = re.sub(r"[^A-Za-z0-9_\-]+", "_", company["name"]).strip("_")
        out_path = os.path.join(out_dir, f"somatie_{safe_name}.docx")
        log(f"[SOMATIE] Writing document: {os.path.basename(out_path)}")
        doc.save(out_path)
        generated.append(out_path)
    return generated


# ----------------------------------------------------------------------
# ENVELOPE (one combined .docx, one page per company)
# ----------------------------------------------------------------------

def generate_envelope_doc(rows, out_dir, log=print):
    log("[ENVELOPE] Reading template structure")
    base_doc = Document(TEMPLATE_ENVELOPE)
    body = base_doc.element.body

    all_paragraph_elements = [p._p for p in base_doc.paragraphs]

    dest_idx = None
    addr_idx = None
    for i, p in enumerate(base_doc.paragraphs):
        text = p.text
        if "Dest:" in text or ("CompanyName" in text):
            if dest_idx is None:
                dest_idx = i
        if "CompanyAddress" in text:
            addr_idx = i

    if dest_idx is None or addr_idx is None:
        raise RuntimeError("Could not locate Dest:/CompanyName/CompanyAddress placeholders in envelope template")

    template_block_elements = [all_paragraph_elements[i] for i in range(dest_idx, addr_idx + 1)]
    last_body_element = all_paragraph_elements[-1]

    prev_last_inserted = last_body_element

    for idx, company in enumerate(rows):
        log(f"[ENVELOPE] Reading company: {company['name']}")
        line1, line2 = format_address_envelope(company["address_raw"])

        insert_after = prev_last_inserted

        cloned_elements = []
        for tmpl_el in template_block_elements:
            new_el = copy.deepcopy(tmpl_el)
            insert_after.addnext(new_el)
            insert_after = new_el
            cloned_elements.append(new_el)

        new_paragraphs = [Paragraph(el, base_doc.paragraphs[0]._parent) for el in cloned_elements]

        for p in new_paragraphs:
            text_now = p.text
            if "CompanyName" in text_now:
                _replace_token_in_paragraph(p, "CompanyName", company["name"].upper())
            elif "CompanyAddress" in text_now:
                if p.runs:
                    bold = p.runs[0].bold
                    for r in list(p.runs):
                        r.text = ""
                    p.runs[0].text = line1
                    p.runs[0].bold = bold
                    br_run = p.add_run()
                    br_run.add_break(WD_BREAK.LINE)
                    line2_run = p.add_run(line2)
                    line2_run.bold = bold
                else:
                    p.add_run(f"{line1}\n{line2}")

        if idx < len(rows) - 1:
            log(f"[ENVELOPE] Writing page break after: {company['name']}")
            page_break_p = base_doc.add_paragraph()
            run = page_break_p.add_run()
            run.add_break(WD_BREAK.PAGE)
            pb_el = page_break_p._p
            body.remove(pb_el)
            cloned_elements[-1].addnext(pb_el)
            prev_last_inserted = pb_el
        else:
            prev_last_inserted = cloned_elements[-1]

    for tmpl_el in template_block_elements:
        tmpl_el.getparent().remove(tmpl_el)

    out_path = os.path.join(out_dir, "envelope_toate_companiile.docx")
    log(f"[ENVELOPE] Writing combined document: {os.path.basename(out_path)}")
    base_doc.save(out_path)
    return out_path


# ----------------------------------------------------------------------
# BORDEROU (one combined .docx, one table row per company)
# ----------------------------------------------------------------------

def _set_cell_text(cell, text):
    first = True
    for paragraph in cell.paragraphs:
        if first:
            if paragraph.runs:
                bold = paragraph.runs[0].bold
                for run in list(paragraph.runs):
                    run.text = ""
                paragraph.runs[0].text = text
                paragraph.runs[0].bold = bold
            else:
                paragraph.add_run(text)
            first = False
        else:
            for run in list(paragraph.runs):
                run.text = ""


def generate_borderou_doc(rows, out_dir, log=print):
    log("[BORDEROU] Reading template table")
    doc = Document(TEMPLATE_BORDEROU)
    table = doc.tables[0]

    template_row_idx = None
    for i, row_ in enumerate(table.rows):
        row_text = " ".join(c.text for c in row_.cells)
        if "CompanyName" in row_text and "CompanyAddress" in row_text:
            template_row_idx = i
            break

    if template_row_idx is None:
        raise RuntimeError("Could not find a CompanyName/CompanyAddress placeholder row in borderou table")

    template_tr = table.rows[template_row_idx]._tr

    total_available_rows = len(table.rows) - 1  # minus header row
    needed = len(rows)

    if needed > total_available_rows:
        for _ in range(needed - total_available_rows):
            new_tr = copy.deepcopy(template_tr)
            table._tbl.append(new_tr)

    for i, company in enumerate(rows):
        log(f"[BORDEROU] Reading company: {company['name']}")
        data_row = table.rows[i + 1]
        address_formatted = format_address_borderou(company["address_raw"])
        _set_cell_text(data_row.cells[1], company["name"])
        _set_cell_text(data_row.cells[2], address_formatted)

    for i in range(needed, total_available_rows):
        data_row = table.rows[i + 1]
        _set_cell_text(data_row.cells[1], "")
        _set_cell_text(data_row.cells[2], "")

    out_path = os.path.join(out_dir, "borderou_toate_companiile.docx")
    log(f"[BORDEROU] Writing combined document: {os.path.basename(out_path)}")
    doc.save(out_path)
    return out_path


# ----------------------------------------------------------------------
# ENTRY POINT
# ----------------------------------------------------------------------

def run_batch(log=print):
    rows = read_source_rows(log=log)
    if not rows:
        log("No companies matched the filter (column L == 'somam'/'som\u0103m'). Nothing to generate.")
        return []

    out_dir = BASE_DIR
    generated_files = []

    generated_files.extend(generate_somatie_docs(rows, out_dir, log=log))
    generated_files.append(generate_envelope_doc(rows, out_dir, log=log))
    generated_files.append(generate_borderou_doc(rows, out_dir, log=log))

    log(f"Done. {len(generated_files)} file(s) generated.")
    return generated_files
